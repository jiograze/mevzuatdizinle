"""
Belge işleme ana sınıfı - PDF, Word, vs. dosyaları işler
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime

# PDF işleme
import PyPDF2
import pdfplumber
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Word işleme  
from docx import Document as DocxDocument

# OCR (opsiyonel)
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from ..utils.text_processor import TextProcessor
from ..utils.document_classifier import DocumentClassifier


class DocumentProcessor:
    """Belge işleme ve analiz sınıfı"""
    
    def __init__(self, config_manager, database_manager):
        self.config = config_manager
        self.db = database_manager
        self.logger = logging.getLogger(self.__class__.__name__)
        
        # Alt bileşenler
        self.text_processor = TextProcessor(config_manager)
        self.document_classifier = DocumentClassifier(config_manager)
        
        # OCR ayarları
        self.ocr_enabled = config_manager.get('ocr.enabled', False)
        self.ocr_confidence_threshold = config_manager.get('ocr.confidence_threshold', 75)
        
        # Desteklenen dosya türleri
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.txt': self._process_txt
        }
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Ana dosya işleme fonksiyonu
        
        Args:
            file_path: İşlenecek dosya yolu
            
        Returns:
            İşlem sonucu dict
        """
        try:
            file_path_obj = Path(file_path)
            self.logger.info(f"Dosya işleme başlatıldı: {file_path}")
            
            # Dosya varlık kontrolü
            if not file_path_obj.exists():
                return {'success': False, 'error': 'Dosya bulunamadı'}
            
            # Dosya boyutu kontrolü
            file_size = file_path_obj.stat().st_size
            max_size = self.config.get('max_file_size_mb', 50) * 1024 * 1024  # MB to bytes
            if file_size > max_size:
                return {'success': False, 'error': f'Dosya çok büyük: {file_size/1024/1024:.1f}MB (max: {max_size/1024/1024}MB)'}
            
            # Dosya türü kontrolü
            file_extension = file_path_obj.suffix.lower()
            if file_extension not in self.supported_extensions:
                return {'success': False, 'error': f'Desteklenmeyen dosya türü: {file_extension}'}
            
            # Dosya hash kontrolü (duplicate check)
            file_hash = self._compute_file_hash(file_path_obj)
            if self._is_duplicate_file(file_hash):
                return {'success': False, 'error': 'Bu dosya daha önce sisteme eklenmiş'}
            
            # Dosya bilgilerini al
            file_info = self._get_file_info(file_path_obj)
            file_info['file_hash'] = file_hash
            
            # Metin çıkarma
            processor_func = self.supported_extensions[file_extension]
            extraction_result = processor_func(file_path_obj)
            
            if not extraction_result['success']:
                return extraction_result
            
            raw_text = extraction_result['text']
            metadata = extraction_result.get('metadata', {})
            
            # Minimum metin uzunluğu kontrolü
            if len(raw_text.strip()) < 50:
                return {'success': False, 'error': 'Dosyada yeterli metin bulunamadı'}
            
            # Belge türü ve bilgileri tespit et
            classification_result = self.document_classifier.classify_document(
                raw_text, 
                file_path_obj.name, 
                metadata
            )
            
            # Maddelere bölme
            articles = self.text_processor.extract_articles(raw_text)
            
            # Eğer hiç madde bulunamadıysa tüm metni tek madde olarak al
            if not articles:
                articles = [{
                    'number': None,
                    'title': None,
                    'content': raw_text.strip(),
                    'type': 'FULL_TEXT'
                }]
            
            # Mülga/değişiklik tespiti
            for article in articles:
                article.update(self.text_processor.detect_amendments(article['content']))
            
            # Veritabanına kaydetme
            save_result = self._save_to_database(
                file_info, 
                classification_result, 
                articles, 
                raw_text, 
                metadata
            )
            
            if save_result['success']:
                # Dosya organizasyonu (eğer etkinse)
                organization_result = None
                if self.config.get('file_organization.enabled', True):
                    organization_result = self.organize_file(
                        str(file_path_obj), 
                        classification_result, 
                        file_info
                    )
                
                self.logger.info(f"Dosya başarıyla işlendi: {file_path}")
                result = {
                    'success': True,
                    'document_id': save_result['document_id'],
                    'articles_count': len(articles),
                    'classification': classification_result,
                    'file_info': file_info,
                    'text_length': len(raw_text)
                }
                
                # Organizasyon bilgilerini ekle
                if organization_result:
                    result['organization'] = organization_result
                
                return result
            else:
                return save_result
                
        except Exception as e:
            self.logger.error(f"Dosya işleme hatası: {file_path} - {e}")
            return {'success': False, 'error': str(e)}
    
    def _get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Dosya bilgilerini topla"""
        stat = file_path.stat()
        
        return {
            'original_filename': file_path.name,
            'file_path': str(file_path.absolute()),
            'file_size': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat()
        }
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """PDF dosyası işleme"""
        try:
            text_content = ""
            metadata = {}
            
            # Önce pdfplumber ile dene
            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata.update({
                        'pages': len(pdf.pages),
                        'pdf_metadata': pdf.metadata or {}
                    })
                    
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
            
            except Exception as e:
                self.logger.warning(f"pdfplumber hatası: {e}, PyPDF2 deneniyor...")
                
                # PyPDF2 ile dene
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        metadata['pages'] = len(pdf_reader.pages)
                        
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + "\n\n"
                            
                except Exception as e2:
                    if PYMUPDF_AVAILABLE:
                        # PyMuPDF ile dene
                        try:
                            doc = fitz.open(file_path)
                            metadata['pages'] = doc.page_count
                            
                            for page in doc:
                                text_content += page.get_text() + "\n\n"
                            
                            doc.close()
                            
                        except Exception as e3:
                            return {'success': False, 'error': f'PDF okuma hatası: {e3}'}
                    else:
                        return {'success': False, 'error': f'PDF okuma hatası: {e2}'}
            
            # OCR kontrolü (metin yoksa)
            if not text_content.strip() and self.ocr_enabled and OCR_AVAILABLE:
                self.logger.info("Metin bulunamadı, OCR deneniyor...")
                ocr_result = self._perform_ocr(file_path)
                if ocr_result['success']:
                    text_content = ocr_result['text']
                    metadata['ocr_used'] = True
            
            if not text_content.strip():
                return {'success': False, 'error': 'PDF\'den metin çıkarılamadı'}
            
            return {
                'success': True,
                'text': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': f'PDF işleme hatası: {e}'}
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """DOCX dosyası işleme"""
        try:
            doc = DocxDocument(file_path)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Tablolar varsa onları da al
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            metadata = {
                'paragraphs_count': len(doc.paragraphs),
                'tables_count': len(doc.tables),
                'core_properties': {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
                }
            }
            
            return {
                'success': True,
                'text': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': f'DOCX işleme hatası: {e}'}
    
    def _process_doc(self, file_path: Path) -> Dict[str, Any]:
        """DOC dosyası işleme (eski format)"""
        # DOC formatı daha karmaşık, python-docx2txt veya antiword kullanılabilir
        # Şimdilik basit bir yaklaşım
        try:
            # Antiword varsa onu kullan
            import subprocess
            result = subprocess.run(
                ['antiword', str(file_path)], 
                capture_output=True, 
                text=True, 
                encoding='utf-8'
            )
            
            if result.returncode == 0:
                return {
                    'success': True,
                    'text': result.stdout,
                    'metadata': {'extraction_method': 'antiword'}
                }
            else:
                return {'success': False, 'error': 'DOC dosyası işlenemedi (antiword gerekli)'}
                
        except FileNotFoundError:
            return {'success': False, 'error': 'DOC dosyası işleme için antiword kurulu değil'}
        except Exception as e:
            return {'success': False, 'error': f'DOC işleme hatası: {e}'}
    
    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """TXT dosyası işleme"""
        try:
            # Encoding algılama
            import chardet
            
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result.get('encoding', 'utf-8')
            
            # Metni oku
            with open(file_path, 'r', encoding=encoding) as f:
                text_content = f.read()
            
            metadata = {
                'encoding': encoding,
                'confidence': encoding_result.get('confidence', 0.0),
                'char_count': len(text_content)
            }
            
            return {
                'success': True,
                'text': text_content,
                'metadata': metadata
            }
            
        except Exception as e:
            return {'success': False, 'error': f'TXT işleme hatası: {e}'}
    
    def _perform_ocr(self, file_path: Path) -> Dict[str, Any]:
        """OCR ile metin çıkarma"""
        if not OCR_AVAILABLE:
            return {'success': False, 'error': 'OCR kütüphaneleri kurulu değil'}
        
        try:
            # PDF'i görsel olarak işle
            if PYMUPDF_AVAILABLE:
                doc = fitz.open(file_path)
                text_content = ""
                
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    # Sayfayı görsel olarak render et
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    # PIL Image'e çevir
                    from io import BytesIO
                    img = Image.open(BytesIO(img_data))
                    
                    # OCR uygula
                    page_text = pytesseract.image_to_string(
                        img, 
                        lang=self.config.get('ocr.lang', 'tur'),
                        config='--oem 3 --psm 6'
                    )
                    
                    # Güven skorunu kontrol et (basit yaklaşım)
                    if len(page_text.strip()) > 10:  # Minimum metin uzunluğu
                        text_content += page_text + "\n\n"
                
                doc.close()
                
                return {
                    'success': True,
                    'text': text_content,
                    'metadata': {'ocr_method': 'tesseract'}
                }
            else:
                return {'success': False, 'error': 'OCR için PyMuPDF gerekli'}
                
        except Exception as e:
            return {'success': False, 'error': f'OCR hatası: {e}'}
    
    def _compute_file_hash(self, file_path: Path) -> str:
        """Dosya MD5 hash'i hesapla"""
        try:
            import hashlib
            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5()
                chunk_size = 8192
                
                while chunk := f.read(chunk_size):
                    file_hash.update(chunk)
                
                return file_hash.hexdigest()
                
        except Exception as e:
            self.logger.error(f"Hash hesaplama hatası: {file_path} - {e}")
            return ""
    
    def _is_duplicate_file(self, file_hash: str) -> bool:
        """Hash'e göre duplicate dosya kontrolü"""
        if not file_hash:
            return False
            
        try:
            cursor = self.db.connection.cursor()
            cursor.execute("""
                SELECT id FROM documents WHERE file_hash = ? LIMIT 1
            """, (file_hash,))
            
            result = cursor.fetchone()
            cursor.close()
            
            return result is not None
            
        except Exception as e:
            self.logger.error(f"Duplicate kontrol hatası: {e}")
            return False
    
    def _save_to_database(self, file_info: Dict, classification: Dict, 
                         articles: List[Dict], raw_text: str, metadata: Dict) -> Dict[str, Any]:
        """İşlenmiş belgeyi veritabanına kaydet"""
        try:
            # Dosya hash'i file_info'dan al (process_file'da zaten hesaplandı)
            file_hash = file_info.get('file_hash', '')
            
            # Belge kaydı
            document_data = {
                'title': classification.get('title', file_info['original_filename']),
                'law_number': classification.get('law_number'),
                'document_type': classification.get('document_type', 'DIGER'),
                'category': classification.get('category'),
                'subcategory': classification.get('subcategory'),
                'original_filename': file_info['original_filename'],
                'stored_filename': self._generate_stored_filename(classification, file_info),
                'file_path': file_info['file_path'],
                'file_hash': file_hash,
                'file_size': file_info['file_size'],
                'effective_date': classification.get('effective_date'),
                'publication_date': classification.get('publication_date'),
                'metadata': str(metadata)  # JSON string olarak
            }
            
            document_id = self.db.insert_document(document_data)
            
            # Madde kayıtları
            for idx, article in enumerate(articles):
                article_data = {
                    'document_id': document_id,
                    'article_number': article.get('number'),
                    'title': article.get('title'),
                    'content': article['content'],
                    'content_clean': self.text_processor.clean_text(article['content']),
                    'seq_index': idx + 1,
                    'is_repealed': article.get('is_repealed', False),
                    'is_amended': article.get('is_amended', False),
                    'amendment_info': article.get('amendment_info'),
                    'article_type': article.get('type', 'MADDE')
                }
                
                self.db.insert_article(article_data)
            
            self.logger.info(f"Veritabanına kaydedildi: doc_id={document_id}, articles={len(articles)}")
            
            return {
                'success': True,
                'document_id': document_id
            }
            
        except Exception as e:
            self.logger.error(f"Veritabanı kayıt hatası: {e}")
            return {'success': False, 'error': f'Veritabanı kayıt hatası: {e}'}
    
    def _generate_stored_filename(self, classification: Dict, file_info: Dict) -> str:
        """Saklanacak dosya adını oluştur"""
        law_number = classification.get('law_number', 'NA')
        title = classification.get('title', file_info['original_filename'])
        doc_type = classification.get('document_type', 'DIGER')
        
        # Türkçe karakterleri normalize et ve slugify
        title_clean = self.text_processor.clean_text(title)
        title_slug = self.text_processor.slugify(title_clean)
        
        # Dosya uzantısı
        original_path = Path(file_info['original_filename'])
        extension = original_path.suffix
        
        # Dosya adı formatı: {numara}_{tip}_{baslik}.{uzanti}
        if law_number and law_number != 'NA':
            # Numara varsa: 2022-4_genelge_dijitallesme.txt
            filename_base = f"{law_number}_{doc_type.lower()}_{title_slug[:30]}"
        else:
            # Numara yoksa: yonetmelik_cevre_koruma.txt  
            filename_base = f"{doc_type.lower()}_{title_slug[:40]}"
        
        return f"{filename_base}{extension}"
    
    def _generate_organized_path(self, classification: Dict, file_info: Dict) -> Path:
        """Organize edilmiş dosya yolunu oluştur"""
        try:
            # Base depo klasörü
            base_folder = self.config.get_base_folder()
            organized_folder = base_folder / 'organized'
            
            # Belge türüne göre ana klasör
            doc_type = classification.get('document_type', 'DIGER')
            doc_type_folder = self._get_document_type_folder(doc_type)
            
            # Yıl ve numara bilgisi
            law_number = classification.get('law_number')
            pub_date = classification.get('publication_date')
            
            # Yıl çıkarma
            year = None
            if pub_date:
                try:
                    year = pub_date.split('-')[0]
                except:
                    pass
            
            # Dosya adından yıl çıkarma (genelge 2022/4 gibi)
            if not year:
                title = classification.get('title', '')
                filename = file_info['original_filename']
                year_match = re.search(r'20\d{2}', f"{title} {filename}")
                if year_match:
                    year = year_match.group()
            
            # Yıl yoksa geçerli yıl
            if not year:
                year = str(datetime.now().year)
            
            # Klasör yapısını oluştur
            target_path = organized_folder / doc_type_folder
            
            # Yıl klasörü ekle
            if year:
                target_path = target_path / year
            
            # Numara klasörü ekle (genelge/yönetmelik için)
            if law_number or doc_type in ['YÖNERGE', 'YÖNETMELİK']:
                # Genelge numarası çıkarma (2022/4 formatı)
                number = law_number
                if not number:
                    # Dosya adından numara çıkarma
                    title = classification.get('title', '')
                    filename = file_info['original_filename']
                    
                    # 2022/4, 2023-12, No:15 gibi formatları ara
                    number_patterns = [
                        r'(\d{4}[/-]\d+)',  # 2022/4, 2023-12
                        r'[Nn]o\s*:?\s*(\d+)',  # No:4, no 15
                        r'sayı\s*:?\s*(\d+)',  # Sayı: 25
                        r'(\d+)\s*[/-]\s*\d{4}',  # 4/2022
                    ]
                    
                    for pattern in number_patterns:
                        match = re.search(pattern, f"{title} {filename}")
                        if match:
                            number = match.group(1)
                            break
                
                if number:
                    target_path = target_path / str(number)
            
            return target_path
            
        except Exception as e:
            self.logger.error(f"Organize yol oluşturma hatası: {e}")
            # Hata durumunda basit yapı
            base_folder = self.config.get_base_folder()
            return base_folder / 'organized' / 'DIGER'
    
    def _get_document_type_folder(self, doc_type: str) -> str:
        """Belge türüne göre klasör adını döndür"""
        folder_mapping = {
            'ANAYASA': 'anayasa',
            'KANUN': 'kanun',
            'KHK': 'khk',
            'TÜZÜK': 'tuzuk',
            'YÖNETMELİK': 'yonetmelik',
            'YÖNERGE': 'genelge',  # Genelge ve yönerge aynı klasörde
            'TEBLİĞ': 'teblig',
            'KARAR': 'karar',
            'DIGER': 'diger'
        }
        
        return folder_mapping.get(doc_type, 'diger')
    
    def organize_file(self, file_path: str, classification: Dict, file_info: Dict) -> Dict[str, Any]:
        """Dosyayı organize et ve uygun klasöre taşı"""
        try:
            source_path = Path(file_path)
            
            # Hedef yolu oluştur
            target_base = self._generate_organized_path(classification, file_info)
            
            # Klasörleri oluştur
            target_base.mkdir(parents=True, exist_ok=True)
            
            # Hedef dosya adı
            stored_filename = self._generate_stored_filename(classification, file_info)
            target_file = target_base / stored_filename
            
            # Eğer aynı isimde dosya varsa numara ekle
            counter = 1
            original_target = target_file
            while target_file.exists():
                name_part = original_target.stem
                extension = original_target.suffix
                target_file = target_base / f"{name_part}_{counter:02d}{extension}"
                counter += 1
            
            # Dosyayı taşı (kopyala + sil)
            import shutil
            shutil.copy2(source_path, target_file)
            
            # Orijinal dosyayı sil (opsiyonel)
            if self.config.get('file_organization.delete_original', True):
                source_path.unlink()
                self.logger.info(f"Orijinal dosya silindi: {source_path}")
            
            self.logger.info(f"Dosya organize edildi: {source_path} -> {target_file}")
            
            return {
                'success': True,
                'target_path': str(target_file),
                'target_folder': str(target_base),
                'organized_structure': self._get_folder_structure(target_base)
            }
            
        except Exception as e:
            self.logger.error(f"Dosya organizasyon hatası: {file_path} - {e}")
            return {
                'success': False,
                'error': f"Dosya organizasyon hatası: {e}"
            }
    
    def _get_folder_structure(self, path: Path) -> str:
        """Klasör yapısını string olarak döndür"""
        try:
            base_folder = self.config.get_base_folder()
            relative_path = path.relative_to(base_folder)
            return str(relative_path).replace('\\', '/')
        except:
            return str(path)
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """İşleme istatistiklerini döndür"""
        # TODO: İmplementasyon
        return {
            'total_processed': 0,
            'success_rate': 0.0,
            'average_processing_time': 0.0
        }
